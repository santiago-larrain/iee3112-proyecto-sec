"""
Extractor de texto de archivos DOCX
"""

from docx import Document
from pathlib import Path
from typing import Optional, Dict, Any
import logging

logger = logging.getLogger(__name__)


class DOCXExtractor:
    """Extrae texto de archivos DOCX usando python-docx"""
    
    def extract_text(self, file_path: Path) -> Optional[str]:
        """
        Extrae texto de un archivo DOCX
        
        Args:
            file_path: Ruta al archivo DOCX
            
        Returns:
            Texto extraído o None si hay error
        """
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extraer texto de párrafos
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extraer texto de tablas
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n\n".join(text_parts) if text_parts else None
            
        except Exception as e:
            logger.error(f"Error extrayendo texto de DOCX {file_path}: {e}")
            return None
    
    def extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """
        Extrae metadatos básicos del DOCX
        
        Args:
            file_path: Ruta al archivo DOCX
            
        Returns:
            Diccionario con metadatos
        """
        metadata = {
            "file_name": file_path.name,
            "file_size": file_path.stat().st_size,
            "num_paragraphs": 0,
            "num_tables": 0
        }
        
        try:
            doc = Document(file_path)
            metadata["num_paragraphs"] = len(doc.paragraphs)
            metadata["num_tables"] = len(doc.tables)
        except Exception as e:
            logger.warning(f"Error obteniendo metadatos de DOCX {file_path}: {e}")
        
        return metadata

